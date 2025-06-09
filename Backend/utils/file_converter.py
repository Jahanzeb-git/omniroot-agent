import os
import logging
from typing import Optional
import tempfile
import subprocess

logger = logging.getLogger(__name__)

def convert_pdf_to_text(pdf_path: str) -> str:
    """Convert PDF to text using pdfplumber for better accuracy."""
    try:
        import pdfplumber
        
        text_content = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- Page {page_num} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
        
        if not text_content:
            logger.error("No text content extracted from PDF")
            return "Error: Could not extract readable text from PDF file."
        
        return "\n\n".join(text_content)
        
    except ImportError:
        logger.error("pdfplumber not installed")
        return "Error: PDF processing library not available."
    except Exception as e:
        logger.error(f"PDF conversion failed: {str(e)}")
        return f"Error: Failed to convert PDF to text: {str(e)}"

def convert_docx_to_text(docx_path: str) -> str:
    """Convert DOCX to text using python-docx."""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text.append(row_text)
            if table_text:
                text_content.append("\n--- Table ---\n" + "\n".join(table_text))
        
        if not text_content:
            logger.error("No text content extracted from DOCX")
            return "Error: Could not extract readable text from DOCX file."
        
        return "\n\n".join(text_content)
        
    except ImportError:
        logger.error("python-docx not installed")
        return "Error: DOCX processing library not available."
    except Exception as e:
        logger.error(f"DOCX conversion failed: {str(e)}")
        return f"Error: Failed to convert DOCX to text: {str(e)}"

def convert_file_to_text(file_path: str) -> Optional[str]:
    """
    Convert PDF or DOCX files to text format.
    Returns converted text or None if file doesn't need conversion.
    """
    if not os.path.exists(file_path):
        return None
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        logger.info(f"Converting PDF to text: {file_path}")
        return convert_pdf_to_text(file_path)
    elif file_ext == '.docx':
        logger.info(f"Converting DOCX to text: {file_path}")
        return convert_docx_to_text(file_path)
    
    # File doesn't need conversion
    return None

def is_convertible_file(file_path: str) -> bool:
    """Check if file is PDF or DOCX that can be converted to text."""
    file_ext = os.path.splitext(file_path)[1].lower()
    return file_ext in ['.pdf', '.docx']